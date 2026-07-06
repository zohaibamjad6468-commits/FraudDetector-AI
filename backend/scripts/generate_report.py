import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()

    # Title Page
    title = doc.add_heading('Final Project Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('FinGuard AI - Credit Card Fraud Intelligence System')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].bold = True
    
    doc.add_paragraph('\n\n')

    # Project Details
    details = doc.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = details.add_run('Subject: Software Verification and Validation (SVV)\n')
    run.bold = True
    run.font.size = Pt(14)
    run = details.add_run('Instructor: Dr. Khurram\n')
    run.bold = True
    run.font.size = Pt(14)
    run = details.add_run('Semester: 6\n\n')
    run.bold = True
    run.font.size = Pt(14)

    # Group Members
    members_heading = doc.add_heading('Project Members:', level=2)
    members_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    members = [
        "1. Muhammad Husnain (Roll No. 232692)",
        "2. Usman Akbar (Roll No. 232670)",
        "3. Muhammad Zohaib (Roll No. 232698)",
        "4. Muhammad Umar Aziz (Roll No. 2326--)"
    ]
    
    for member in members:
        p = doc.add_paragraph(member)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(12)

    doc.add_page_break()

    # 1. Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        "FinGuard AI is an advanced, machine-learning-driven Credit Card Fraud Detection system. "
        "The project aims to provide financial institutions with a robust tool to automatically analyze "
        "transactions, detect anomalies, and flag potentially fraudulent activities in real-time."
    )

    # 2. Project Architecture
    doc.add_heading('2. System Architecture', level=1)
    doc.add_paragraph(
        "The system is built using a modern 3-tier architecture to ensure scalability, "
        "security, and maintainability:"
    )
    doc.add_paragraph("Frontend: Built with React.js and Vite, featuring a responsive Glassmorphism UI.", style='List Bullet')
    doc.add_paragraph("Backend: Developed in Python using the Flask framework. It follows a Service-Layer pattern to separate business logic from routing.", style='List Bullet')
    doc.add_paragraph("Database: MySQL 8.0 with a fully normalized 10-table schema to maintain data integrity.", style='List Bullet')
    doc.add_paragraph("Machine Learning: Utilizes XGBoost, trained on the Kaggle ULB dataset, with SMOTE applied for handling class imbalance.", style='List Bullet')

    # 3. Core Features
    doc.add_heading('3. Core Features Implemented', level=1)
    doc.add_paragraph("1. Real-Time Fraud Detection: Integrates a trained XGBoost model for sub-millisecond predictions.", style='List Number')
    doc.add_paragraph("2. Synthetic Feature Generation: Dynamically generates V1-V28 PCA features from business-level inputs (Merchant, Amount, Location).", style='List Number')
    doc.add_paragraph("3. Explainable AI (XAI): Displays the top 3 factors contributing to a fraud flag.", style='List Number')
    doc.add_paragraph("4. Role-Based Access Control (RBAC): Differentiates between 'Admin' and 'Analyst' roles via JWT.", style='List Number')
    doc.add_paragraph("5. Live Toast Notifications: Asynchronous polling alerts analysts immediately when a high-risk case arrives.", style='List Number')
    doc.add_paragraph("6. Session Management (FR-03): Automatically logs users out after 15 minutes of inactivity for security.", style='List Number')
    doc.add_paragraph("7. Compliance & Exporting: Immutable Audit Logs, and PDF/CSV export functionality for transaction reporting.", style='List Number')

    doc.add_page_break()

    # 4. Screenshots
    doc.add_heading('4. Application Interface (Screenshots)', level=1)
    
    screenshots_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', 'docs', 'screenshots'))
    
    if os.path.exists(screenshots_dir):
        files = sorted([f for f in os.listdir(screenshots_dir) if f.endswith(('.png', '.jpg', '.jpeg'))])
        for idx, filename in enumerate(files):
            doc.add_heading(f"Figure {idx + 1}: Interface View", level=3)
            img_path = os.path.join(screenshots_dir, filename)
            try:
                doc.add_picture(img_path, width=Inches(6.0))
            except Exception as e:
                doc.add_paragraph(f"[Could not load image: {filename} - {str(e)}]")
            doc.add_paragraph("\n")
    else:
        doc.add_paragraph("No screenshots directory found.")

    # 5. Conclusion
    doc.add_page_break()
    doc.add_heading('5. Conclusion', level=1)
    doc.add_paragraph(
        "The FinGuard AI project successfully implements all the requirements outlined in the initial proposal. "
        "By utilizing modern web technologies, strict security measures (JWT, RBAC, Auto-logout), and an optimized "
        "XGBoost machine learning pipeline, the system serves as a highly capable and professional-grade solution "
        "for detecting financial fraud."
    )

    # Save Document
    output_path = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', 'FinGuard_AI_Final_Report.docx'))
    doc.save(output_path)
    print(f"Report successfully generated at: {output_path}")

if __name__ == "__main__":
    create_report()
